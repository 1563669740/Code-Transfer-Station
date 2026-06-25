"""生成导出 DOCX 文件：Codex 远程执行系统部署说明。

注意：本脚本只生成文档，不运行项目业务代码。文档内容必须遵守
`docs/RUNBOOK.md`：Codex 本地禁止执行 run.sh、pytest、make run、make test。
"""
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

PROJECT_DIR = Path(__file__).resolve().parents[1]
OUTPUT = PROJECT_DIR / "Codex远程执行系统-部署全流程记录.docx"


def set_normal_style(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    style = doc.styles["Normal"]
    style.font.name = "微软雅黑"
    style.font.size = Pt(11)


def add_heading(doc: Document, text: str, level: int = 1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    return heading


def add_para(doc: Document, text: str, bold: bool = False):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = bold
    return paragraph


def add_code(doc: Document, text: str):
    for line in text.strip("\n").split("\n"):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Cm(1)
        run = paragraph.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)


def build() -> Path:
    doc = Document()
    set_normal_style(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Codex 远程执行系统\n部署全流程记录")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("本地静态修改 -> GitHub 中转 -> 控制服务器自动拉取执行")
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_page_break()

    add_heading(doc, "1. 核心规则")
    add_para(doc, "Codex 本地只做代码修改、文档修改、静态检查、清单一致性维护和 Git 提交推送。")
    add_para(doc, "Codex 本地禁止运行项目业务代码、run.sh、pytest、make run、make test。", bold=True)
    add_para(doc, "控制服务器拉取新 commit 后，才执行统一入口和测试。")
    add_code(doc, """
# Codex 本地允许
rg "关键字"
git status
git diff
git add -A && git commit -m "描述修改" && git push

# Codex 本地禁止
bash run.sh
python3 -m pytest -q
pytest
make run
make test
""")

    add_heading(doc, "2. Git 拉取模型")
    add_code(doc, """
Windows / Codex（开发机）              控制服务器
  1. 修改代码                             1. 定时 git fetch
  2. 静态检查/清单检查                     2. 检测到新 commit -> git merge --ff-only
  3. git commit && git push               3. bash run.sh
         |                                4. python3 -m pytest -q
         v                                5. 写入 ~/codex_pull_logs/latest.log
    GitHub 仓库  <----------------------- 日志回传可选
""")

    add_heading(doc, "3. Codex 提示词")
    add_code(doc, """
请先阅读 docs/RUNBOOK.md，按其中规则修改代码。
生成或修改代码后，禁止在本地执行 bash run.sh、python3 -m pytest -q、pytest、make run 或 make test。
如果涉及新的核心业务入口，请同步更新 docs/CURRENT_BUSINESS.json、main.py、src/ 和 tests/，删除旧业务残留。
只做不执行业务代码的静态检查和清单检查，然后 git add -A、git commit、git push；由控制服务器自动拉取后执行 run.sh 和 pytest。
""")

    add_heading(doc, "4. 服务器没有自动更新时")
    add_para(doc, "逻辑代码已经 push 但服务器没有自动更新时，由用户在控制服务器上执行同步脚本；Codex 不通过 SSH 代替执行。")
    add_code(doc, """
cd ~/codex_projects/project
bash scripts/server_sync_latest.sh
bash scripts/server_sync_latest.sh --no-run-once
bash scripts/server_sync_latest.sh --no-restart-daemon
""")

    add_heading(doc, "5. 外部数据")
    add_para(doc, "如果业务代码需要项目目录外的数据，不要写死本机路径，也不要提交大数据或隐私数据。")
    add_code(doc, """
# 服务器上准备数据目录，并用环境变量传入
DATA_DIR=/data/code-transfer-station/input

# 运行产物继续交给服务器脚本归档
RUN_ARTIFACT_DIR=~/codex_pull_logs/artifacts/latest
""")

    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(f"已生成：{build()}")
