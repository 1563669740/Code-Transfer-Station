"""生成导出DOCX文件：Codex本地与服务器配置指南"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

PROJECT_DIR = os.path.dirname(os.path.abspath(__file__))

doc = Document()

# ── 页面设置 ──
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ── 样式 ──
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(11)

# ── 辅助函数 ──
def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    return h

def add_para(text, bold=False, italic=False, size=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    return p

def add_code(text):
    """添加代码块"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    for line in text.strip().split('\n'):
        run = p.add_run(line + '\n')
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
    return p

def add_cmd(text):
    """添加命令块（深色背景用灰色文字）"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    for line in text.strip().split('\n'):
        run = p.add_run(line + '\n')
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x00, 0x66, 0x33)
    return p

def add_output(text):
    """添加输出块"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    for line in text.strip().split('\n'):
        run = p.add_run(line + '\n')
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    return p

def add_note(text):
    """添加提示"""
    p = doc.add_paragraph()
    run = p.add_run('💡 ' + text)
    run.font.size = Pt(10)
    run.italic = True
    return p

def add_separator():
    doc.add_paragraph('─' * 50)

# ============================================================
# 封面
# ============================================================
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Codex 远程执行系统\n部署全流程记录')
run.bold = True
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('本地开发 → GitHub 中转 → 服务器自动拉取执行')
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run('日期：2024年6月24日\n').font.size = Pt(11)
meta.add_run('仓库：github.com/1563669740/Code-Transfer-Station\n').font.size = Pt(11)
meta.add_run('分支：main').font.size = Pt(11)

doc.add_page_break()

# ============================================================
# 目录
# ============================================================
add_heading('目录', level=1)
toc_items = [
    '1. 项目概述与架构',
    '2. Git 拉取方案说明',
    '3. 本地端准备（Windows / Codex）',
    '4. 本地测试与问题修复',
    '5. GitHub 仓库配置',
    '6. 代码推送',
    '7. 控制服务器配置',
    '8. 服务器测试与守护进程启动',
    '9. 全链路验证总结',
    '10. 后续使用说明',
    '附录：关键命令速查',
]
for item in toc_items:
    add_para(item)

doc.add_page_break()

# ============================================================
# 1. 项目概述
# ============================================================
add_heading('1. 项目概述与架构', level=1)

add_para('本项目实现了一套"本地开发 → GitHub 中转 → 服务器自动拉取执行"的自动化流水线。核心思想是：开发者在本地修改代码并 git push，控制服务器定时轮询 GitHub 仓库，发现新 commit 后自动拉取、执行、记录日志。')

add_heading('架构图', level=2)
add_code("""Windows / Codex（开发机）              控制服务器（Linux）
  1. 修改代码                             1. 定时 git fetch（每60秒）
  2. bash run.sh                          2. 检测到新 commit → git merge --ff-only
  3. python3 -m pytest -q                 3. bash run.sh
  4. git commit && git push               4. python3 -m pytest -q
         |                                5. 写入 ~/codex_pull_logs/latest.log
         v                                       ^
    GitHub 仓库 ─────────────────────────────────┘
""")

add_heading('三条核心命令（不可绕过）', level=2)
add_cmd("""# 本地运行（修改代码后必须先跑通）
bash run.sh

# 本地测试（必须全部通过）
python3 -m pytest -q

# 推送到远端（本地验证通过后提交推送）
git add -A && git commit -m "描述修改内容" && git push
""")

doc.add_page_break()

# ============================================================
# 2. Git 拉取方案说明
# ============================================================
add_heading('2. Git 拉取方案说明', level=1)

add_para('关于本地和服务器是否需要用同一个账号的问题，有以下几种方案：')

add_heading('方案对比', level=2)

# 表格
table = doc.add_table(rows=5, cols=4, style='Table Grid')
headers = ['方案', '服务器需要账号？', '适用场景', '安全性']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for run in p.runs:
            run.bold = True

data = [
    ['同一账号', '是', '个人项目，最简单', '中'],
    ['独立账号 + 协作者', '是', '想隔离身份、控制权限', '高'],
    ['Deploy Key（SSH）', '否', '生产环境，最推荐 ✓', '最高'],
    ['PAT + HTTPS', '否', '不想配 SSH 时', '中'],
]
for row_idx, row_data in enumerate(data, start=1):
    for col_idx, val in enumerate(row_data):
        table.rows[row_idx].cells[col_idx].text = val

add_para('')
add_note('本项目最终采用 Deploy Key 方案：服务器不需要 GitHub 账号，只需要一把只读 SSH 密钥即可自动拉取代码。')

doc.add_page_break()

# ============================================================
# 3. 本地端准备
# ============================================================
add_heading('3. 本地端准备（Windows / Codex）', level=1)

add_heading('3.1 项目初始结构', level=2)
add_code("""project/
  main.py                   项目主入口（只做入口，不放业务逻辑）
  src/                      业务代码
  tests/                    测试代码
  scripts/
    server_pull_run.sh      控制服务器常驻轮询守护进程
    server_pull_once.sh     控制服务器 cron 单次轮询
    server_pull_log.sh      控制服务器日志回传（可选）
  run.sh                    本地统一启动脚本
  requirements.txt          Python 依赖
  RUNBOOK.md                Codex 执行规则（完整版）
  .gitignore
""")

add_heading('3.2 创建业务代码：src/test.py', level=2)
add_code("""def hello():
    return "hello world"


if __name__ == "__main__":
    print(hello())
""")

add_heading('3.3 修改入口：main.py', level=2)
add_code("""from src.test import hello


def main():
    print(hello())


if __name__ == "__main__":
    main()
""")

add_heading('3.4 修改测试：tests/test_smoke.py', level=2)
add_code("""from src.test import hello


def test_hello():
    assert hello() == "hello world"
""")

doc.add_page_break()

# ============================================================
# 4. 本地测试与问题修复
# ============================================================
add_heading('4. 本地测试与问题修复', level=1)

add_heading('4.1 初始测试', level=2)
add_cmd('$ python main.py')
add_output('hello world')

add_cmd('$ python -m pytest -q')
add_output('.                                                    [100%]\n1 passed in 0.02s')

add_heading('4.2 发现的问题及修复', level=2)

add_para('问题 1：没有配置 git remote', bold=True)
add_output('$ git remote -v\n（空）')
add_para('修复：')
add_cmd('$ git remote add origin https://github.com/1563669740/Code-Transfer-Station.git')

add_para('')
add_para('问题 2：分支名不匹配', bold=True)
add_para('本地分支名为 master，服务器脚本默认跟踪 main。')
add_para('修复：')
add_cmd('$ git branch -M main')

add_para('')
add_para('问题 3：python3 不可用（Windows Store 占位）', bold=True)
add_para('Windows 上 python3 指向 Windows Store 假程序，无法运行。实际可用命令为 python。')
add_para('修复 run.sh：')
add_code("""# 优先使用 python，其次 python3
if command -v python >/dev/null 2>&1; then
  python main.py "$@"
elif command -v python3 >/dev/null 2>&1; then
  python3 main.py "$@"
else
  echo "[ERROR] Python not found" >&2
  exit 1
fi
""")

add_para('')
add_para('问题 4：.venv 虚拟环境损坏', bold=True)
add_para('修复：删除损坏的 .venv，直接使用系统 Python。')
add_cmd('$ rm -rf .venv')

add_heading('4.3 修复后验证', level=2)
add_cmd('$ bash run.sh')
add_output('hello world')
add_cmd('$ python -m pytest -q')
add_output('1 passed in 0.02s')

doc.add_page_break()

# ============================================================
# 5. GitHub 仓库配置
# ============================================================
add_heading('5. GitHub 仓库配置', level=1)

add_para('仓库地址：https://github.com/1563669740/Code-Transfer-Station.git')
add_para('仓库类型：私有（Private）')

add_heading('5.1 本地配置 Remote', level=2)
add_cmd('$ git remote add origin https://github.com/1563669740/Code-Transfer-Station.git')

add_heading('5.2 获取 Personal Access Token', level=2)
add_para('1. 打开 GitHub → 右上角头像 → Settings')
add_para('2. 左侧最下面 → Developer settings → Personal access tokens → Tokens (classic)')
add_para('3. Generate new token (classic) → 勾选 repo → 生成')
add_para('4. 复制生成的 token（格式：ghp_xxxxxxxxxxxx）')

add_note('Token 不要写入代码或配置文件。push 时作为密码输入，可用 Git credential 缓存。')

doc.add_page_break()

# ============================================================
# 6. 代码推送
# ============================================================
add_heading('6. 代码推送', level=1)

add_heading('6.1 Push 到 GitHub', level=2)
add_cmd("""$ git add -A
$ git commit -m "添加 hello world 业务代码：src/test.py，修复 run.sh Python 兼容性，修复分支名为 main"
$ git push -u origin main
""")
add_output("""[main d0ebcdc] 添加 hello world 业务代码...
 4 files changed, 24 insertions(+), 4 deletions(-)
 create mode 100644 src/test.py
 To https://github.com/1563669740/Code-Transfer-Station.git
 * [new branch]      main -> main
""")

add_heading('6.2 安全清理', level=2)
add_para('推送完成后立即从 Git 配置中清除 Token：')
add_cmd('$ git remote set-url origin https://github.com/1563669740/Code-Transfer-Station.git')

add_note('token 切勿留在 git remote URL 或 .git/config 中。')

doc.add_page_break()

# ============================================================
# 7. 控制服务器配置
# ============================================================
add_heading('7. 控制服务器配置', level=1)

add_heading('7.1 服务器环境', level=2)
add_code("""OS: Linux (Ubuntu)
Kernel: 5.10.0
Python: 3.10.12
Shell: bash
Username: batchcom
""")

add_heading('7.2 安装 Git', level=2)
add_cmd('$ sudo apt-get install -y git')
add_output('git version 2.34.1')

add_heading('7.3 生成 Deploy Key（SSH 密钥）', level=2)
add_cmd("""$ mkdir -p ~/.ssh && chmod 700 ~/.ssh
$ ssh-keygen -t ed25519 -f ~/.ssh/codex_pull_deploy_key -C "codex-pull-deploy" -N ""
""")
add_output("""ssh-ed25519 AAAAC3NzaC1lZDI1NTE5AAAAILARXPR9CUM2CXR992Amt7xQ303y2Wa4feO9j1ma4iDu codex-pull-deploy
""")

add_heading('7.4 在 GitHub 添加 Deploy Key', level=2)
add_para('1. 打开 https://github.com/1563669740/Code-Transfer-Station/settings/keys')
add_para('2. 点 Add deploy key')
add_para('3. Title 填：server-deploy')
add_para('4. Key 粘贴公钥内容')
add_para('5. 不勾 Allow write access（只读更安全）')
add_para('6. 点 Add key')

add_heading('7.5 配置 SSH Host', level=2)
add_code("""Host github-codex-pull
    HostName github.com
    User git
    IdentityFile ~/.ssh/codex_pull_deploy_key
    IdentitiesOnly yes
""")

add_heading('7.6 测试连接', level=2)
add_cmd('$ ssh -T git@github-codex-pull')
add_output('Hi 1563669740/Code-Transfer-Station! You\'ve successfully authenticated, but GitHub does not provide shell access.')

doc.add_page_break()

# ============================================================
# 8. 服务器测试与守护进程启动
# ============================================================
add_heading('8. 服务器测试与守护进程启动', level=1)

add_heading('8.1 克隆项目', level=2)
add_cmd("""$ mkdir -p ~/codex_projects
$ git clone git@github-codex-pull:1563669740/Code-Transfer-Station.git ~/codex_projects/project
""")

add_heading('8.2 手动验证', level=2)
add_cmd('$ cd ~/codex_projects/project && bash run.sh')
add_output('hello world')

add_para('安装测试依赖：')
add_cmd("""$ sudo apt-get install -y python3-pip
$ pip3 install -r requirements.txt
$ python3 -m pytest -q
""")
add_output('1 passed in 0.01s')

add_heading('8.3 启动常驻守护进程', level=2)
add_cmd("""$ chmod +x scripts/server_pull_run.sh
$ mkdir -p ~/codex_pull_logs
$ nohup bash scripts/server_pull_run.sh > ~/codex_pull_logs/daemon.out 2>&1 &
$ echo $! > ~/codex_pull_logs/daemon.pid
$ ps aux | grep server_pull_run | grep -v grep
""")
add_output('batchcom  7959  ...  bash scripts/server_pull_run.sh')

add_para('守护进程行为：', bold=True)
add_para('1. 每 60 秒 git fetch origin main')
add_para('2. 比较 remote_sha vs last_run_sha')
add_para('3. 发现新 commit → git merge --ff-only → bash run.sh → pytest')
add_para('4. 结果写入 ~/codex_pull_logs/latest.log')
add_para('5. 循环')

add_heading('8.4 查看日志', level=2)
add_cmd("""# 查看最新执行日志
$ tail -n 160 ~/codex_pull_logs/latest.log

# 查看守护进程是否存活
$ ps aux | grep server_pull_run | grep -v grep

# 查看历史日志
$ ls -lh ~/codex_pull_logs/
""")

add_heading('8.5 停止守护进程', level=2)
add_cmd("""$ kill $(cat ~/codex_pull_logs/daemon.pid)
$ rm ~/codex_pull_logs/daemon.pid
""")

doc.add_page_break()

# ============================================================
# 9. 全链路验证总结
# ============================================================
add_heading('9. 全链路验证总结', level=1)

add_heading('9.1 最终文件清单', level=2)

table2 = doc.add_table(rows=12, cols=3, style='Table Grid')
t2_headers = ['文件', '路径', '说明']
for i, h in enumerate(t2_headers):
    table2.rows[0].cells[i].text = h
    for p in table2.rows[0].cells[i].paragraphs:
        for run in p.runs:
            run.bold = True

files_data = [
    ['main.py', '/', '项目入口，调用 src.test.hello()'],
    ['src/test.py', 'src/', '业务代码，hello() 返回 "hello world"'],
    ['tests/test_smoke.py', 'tests/', '测试用例，验证 hello() 返回值'],
    ['run.sh', '/', '本地启动脚本，自动检测 Python'],
    ['requirements.txt', '/', '依赖：pytest'],
    ['RUNBOOK.md', '/', '项目执行规则（Codex 必读）'],
    ['.gitignore', '/', '排除 .env/venv/pytest_cache 等'],
    ['server_pull_run.sh', 'scripts/', '服务器守5进程轮询脚本'],
    ['server_pull_once.sh', 'scripts/', '服务器 cron 单次轮询脚本'],
    ['server_push_log.sh', 'scripts/', '日志回传脚本（可选）'],
    ['config（服务器）', '~/.ssh/', 'SSH Host 配置 + Deploy Key'],
]
for row_idx, row_data in enumerate(files_data, start=1):
    for col_idx, val in enumerate(row_data):
        table2.rows[row_idx].cells[col_idx].text = val

add_heading('9.2 全链路状态', level=2)

table3 = doc.add_table(rows=10, cols=3, style='Table Grid')
t3_headers = ['环节', '状态', '说明']
for i, h in enumerate(t3_headers):
    table3.rows[0].cells[i].text = h
    for p in table3.rows[0].cells[i].paragraphs:
        for run in p.runs:
            run.bold = True

status_data = [
    ['本地 bash run.sh', '✅', '输出 hello world'],
    ['本地 pytest', '✅', '1 passed'],
    ['git remote 配置', '✅', 'HTTPS + Token 认证'],
    ['分支名统一', '✅', '本地和远程均为 main'],
    ['Python 兼容性', '✅', 'run.sh 优先 python，后备 python3'],
    ['GitHub Deploy Key', '✅', 'SSH 只读密钥'],
    ['服务器 clone', '✅', 'git clone 成功'],
    ['服务器手动执行', '✅', 'run.sh + pytest 均通过'],
    ['守护进程', '✅', '运行中，每 60s 轮询'],
]
for row_idx, row_data in enumerate(status_data, start=1):
    for col_idx, val in enumerate(row_data):
        table3.rows[row_idx].cells[col_idx].text = val

doc.add_page_break()

add_heading('9.3 完整自动化流程', level=2)
add_code("""        你的电脑                          控制服务器
           │                                  │
     1. 修改代码                               │
     2. bash run.sh                            │
     3. python -m pytest -q                    │
     4. git push ──────▶ GitHub ◀──── 5. git fetch（每60秒）
                                         6. 发现新commit → merge
                                         7. bash run.sh
                                         8. pytest
                                         9. 写日志 → sleep 60s
""")

# ============================================================
# 10. 后续使用说明
# ============================================================
add_heading('10. 后续使用说明', level=1)

add_heading('10.1 日常开发流程', level=2)
add_para('每次修改代码后：')
add_cmd("""# 1. 修改 src/ 或 tests/ 中的代码

# 2. 本地验证
bash run.sh
python -m pytest -q

# 3. 推送
git add -A && git commit -m "描述修改内容" && git push

# 4. 服务器会在 60 秒内自动拉取执行，无需手动操作
""")

add_heading('10.2 查看远程执行结果', level=2)
add_cmd("""# 在服务器上查看
tail -n 160 ~/codex_pull_logs/latest.log

# 或在本地拉取 run-logs 分支（如果配置了日志回传）
git fetch origin run-logs
git show origin/run-logs:logs/run_*.log | tail -n 80
""")

add_heading('10.3 注意事项', level=2)
add_para('❌ 不要手动 scp / ssh 登录服务器改代码')
add_para('❌ 不要把密码、私钥、Token 写入代码或日志')
add_para('❌ 不要绕过 run.sh 直接猜测启动命令')
add_para('❌ 不要在控制服务器上手改代码（应通过 git push 触发自动拉取）')

doc.add_page_break()

# ============================================================
# 附录
# ============================================================
add_heading('附录：关键命令速查', level=1)

add_heading('本地（Windows / Codex）', level=2)
add_cmd("""# 运行
bash run.sh

# 测试
python -m pytest -q

# 推送
git add -A && git commit -m "描述" && git push

# 查看状态
git status
git remote -v
""")

add_heading('服务器', level=2)
add_cmd("""# 查看守护进程
ps aux | grep server_pull_run | grep -v grep

# 查看最新日志
tail -n 160 ~/codex_pull_logs/latest.log

# 停止守护进程
kill $(cat ~/codex_pull_logs/daemon.pid)

# 重启守护进程
cd ~/codex_projects/project
nohup bash scripts/server_pull_run.sh > ~/codex_pull_logs/daemon.out 2>&1 &

# 手动测试
cd ~/codex_projects/project && bash run.sh
cd ~/codex_projects/project && python3 -m pytest -q

# 测试 SSH 连接
ssh -T git@github-codex-pull
""")

add_heading('GitHub', level=2)
add_cmd("""# 仓库地址
https://github.com/1563669740/Code-Transfer-Station

# Deploy Keys 管理
https://github.com/1563669740/Code-Transfer-Station/settings/keys

# Personal Access Tokens
https://github.com/settings/tokens
""")

# ============================================================
# 保存
# ============================================================
output_path = os.path.join(PROJECT_DIR, 'Codex远程执行系统-部署全流程记录.docx')
doc.save(output_path)
print(f'已生成：{output_path}')
